#!/usr/bin/env python3
"""level0_verify.py — TASK-OBS-0054-бис, ступень 0 ПО СУЩЕСТВУ.

level0_check.py умеет только "есть ли файл с правдоподобным именем" — этого
хватало, пока заявки называли источники словами ("исходящее письмо РСУ-8 на
уведомление о расторжении"), а файлы названы шифрами
(RSU8-OUT-2026-06-29-29-3-06-26.pdf): совпадений по номеру/дате в самом
ИМЕНИ файла почти нет, 16 заявок первого прогона закрылись «проверить не
смог — первоисточник не найден на уровне 0» именно поэтому.

Этот модуль читает СОДЕРЖИМОЕ (текстовый слой, без OCR/сети/моделей — только
то, что уже лежит на диске) и сверяет утверждение заявки с текстом источника
по существу: подтверждено / опровергнуто / не смог — с причиной, годной как
вход для ступени 1.

Не трогает: ladder.py, level0_check.py, selftest.py, executors.py, worker.py
(правят другие воркеры параллельно). Не пишет НИЧЕГО в CSV реестров волта —
`iter_claims()` только читает.

Запуск:
    python level0_verify.py --dry-run     # сводка + построчный отчёт в JSON
    python level0_verify.py --selftest    # assert-самопроверка
"""
from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import sys
import zipfile
from pathlib import Path

_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

# Переиспользуем извлечение номеров/дат и индекс 05_ORIGINALS из level0_check —
# не копипаста, тот же код, что уже прошёл его selftest.
from level0_check import (  # noqa: E402
    ORIGINALS_DIR,
    REGISTRY_DIR,
    REGISTRY_GLOB,
    extract_dates,
    extract_numbers,
    file_signals,
    load_originals,
    split_sources,
)
import ladder  # noqa: E402  — verify_principle_pattern переиспользуем, не дублируем

CACHE_DIR = Path(r"E:\-4-\air-worker\level0_text_cache")
REPORT_PATH = Path(r"E:\-4-\air-worker\level0_verify_report.json")

# Хотя бы один сигнал должен набрать этот балл, иначе кандидата в файле не
# считаем — «нашли хоть что-то похожее» не равно «нашли источник».
MIN_SCORE = 2


# --- 1. извлечение текстового слоя, $0, без сети и без моделей ----------------
#
# Порядок дешевизны: готовый .md/.ocr.txt рядом с файлом (9 штук уже есть) →
# .xlsx/.docx через уже установленные библиотеки → .pdf через pypdf
# (TASK-OBS-0054-бис часть A: своего извлекателя текста в контуре нет —
# air-storage/inside.py умеет читать PDF через pdfminer, но его SKILL.md
# прямо запрещает отдавать текст наружу, только координаты — "содержимое
# документов не поднимается ни в отчёт, ни в лог", это архитектурная
# граница, а не пробел; поэтому `python -m pip install pypdf`, разрешено
# постановкой). pdfplumber — запасной путь, если pypdf вдруг недоступен.
# Скан без текстового слоя (PDF без ToUnicode/чистый растр) честно даёт
# пустую строку — OCR здесь намеренно не запускаем (пункт 2 ТЗ).

def _sidecar_text(path: Path) -> str | None:
    for suf in (".md", ".ocr.txt"):
        sc = Path(str(path) + suf)
        if sc.is_file():
            try:
                return sc.read_text(encoding="utf-8", errors="replace")
            except OSError:
                return None
    return None


def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return ""
    try:
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        chunks = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        chunks.append(str(cell))
        wb.close()
        return " ".join(chunks)
    except Exception:
        return ""  # повреждённый/нестандартный xlsx — законный пустой результат


def _read_docx_zip(path: Path) -> str:
    """python-docx не установлен — разбор word/document.xml из zip стандартной
    библиотекой (zipfile+re), без новой зависимости."""
    try:
        with zipfile.ZipFile(str(path)) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    except Exception:
        return ""
    return re.sub(r"<[^>]+>", " ", xml)


def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError:
        return _read_docx_zip(path)
    try:
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return _read_docx_zip(path)


def _read_pdf(path: Path) -> str:
    """pypdf установлен (часть A TASK-OBS-0054-бис). Скан без текстового
    слоя всё равно честно даёт пустую строку — OCR не зовём (пункт 2 ТЗ)."""
    try:
        import pypdf
    except ImportError:
        pypdf = None
    if pypdf is not None:
        try:
            reader = pypdf.PdfReader(str(path))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception:
            return ""
    try:
        import pdfplumber
    except ImportError:
        return ""
    try:
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    except Exception:
        return ""


def _cache_path(path: Path) -> Path:
    h = hashlib.sha1(str(path.resolve()).encode("utf-8", errors="replace")).hexdigest()
    return CACHE_DIR / f"{h}.txt"


def extract_text(path: str | Path) -> str:
    """Текстовый слой первоисточника, $0. Кэш на диске (не в волте) — текст
    извлекается один раз на файл, ключ кэша учитывает mtime+size файла."""
    path = Path(path)

    sidecar = _sidecar_text(path)
    if sidecar is not None:
        return sidecar  # самый дешёвый путь — уже готовый слой, кэш не нужен

    try:
        stat = path.stat()
    except OSError:
        return ""

    cache_file = _cache_path(path)
    if cache_file.is_file():
        try:
            raw = cache_file.read_text(encoding="utf-8")
            header, _, body = raw.partition("\n")
            mtime_s, size_s = header.split("\t")
            if float(mtime_s) == stat.st_mtime and int(size_s) == stat.st_size:
                return body
        except (OSError, ValueError):
            pass  # кэш повреждён/устарел — извлекаем заново

    suffix = path.suffix.lower()
    if suffix in (".xlsx", ".xlsm"):
        text = _read_xlsx(path)
    elif suffix == ".xls":
        text = ""  # ponytail: xlrd не установлен, старый .xls не читаем — законное "не смог"
    elif suffix == ".docx":
        text = _read_docx(path)
    elif suffix == ".pdf":
        text = _read_pdf(path)
    else:
        text = ""  # .doc/.zip/.gsfx — не поддерживаются (пункт 1 ТЗ)

    try:
        CACHE_DIR.mkdir(parents=True, exist_ok=True)
        cache_file.write_text(f"{stat.st_mtime}\t{stat.st_size}\n{text}", encoding="utf-8")
    except OSError:
        pass  # кэш — оптимизация, не критический путь
    return text


# --- 2. ранжирование кандидатов -----------------------------------------------
#
# Таблица транслитерации сторон и типов документов — по факту того, что
# встречается в реестрах ЦКБА и в шифрах имён файлов 05_ORIGINALS.

PARTY_ALIASES: dict[str, list[str]] = {
    "рсу-8": ["rsu8", "rsu-8"],
    "рсу8": ["rsu8", "rsu-8"],
    "цкба": ["ckba"],
    "молчановагрупп": ["molchanovagrupp"],
    "молчанова групп": ["molchanovagrupp"],
    "псб": ["psb"],
    "лаборатория": ["laba"],
    "приямок": ["priyamok"],
}

DOC_TYPE_ALIASES: dict[str, list[str]] = {
    "письмо": ["pismo", "out", "in", "draft"],
    "исходящ": ["out", "draft"],
    "входящ": ["in"],
    "уведомлен": ["out", "in", "pismo"],
    "договор": ["dogovor"],
    "протокол": ["protokol"],
    "акт": ["act"],
    "смета": ["smeta"],
    "упд": ["упд"],
    "кс-2": ["кс-2", "ks-2"],
    "кс-3": ["кс-3", "ks-3"],
}


def _alias_hits(text_lower: str, table: dict[str, list[str]]) -> set[str]:
    hits: set[str] = set()
    for ru, aliases in table.items():
        if ru in text_lower or any(a in text_lower for a in aliases):
            hits.add(ru)
            hits.update(aliases)
    return hits


def _party_hits(text_lower: str) -> set[str]:
    return _alias_hits(text_lower, PARTY_ALIASES)


def _doctype_hits(text_lower: str) -> set[str]:
    return _alias_hits(text_lower, DOC_TYPE_ALIASES)


def rank_candidates(
    required_source_text: str,
    file_index: list[tuple[Path, set[str], set[str]]],
) -> list[tuple[Path, int, str]]:
    """Ранжированный список кандидатов вместо «первого совпадения».

    Сигналы (баллы): номер документа в имени файла (5), номер договора,
    найденный в имени ПАПКИ — SMETA-5239-47 ↔ №5239/47 (4), дата (3), сторона
    по таблице транслитерации (2), тип документа по ключевым словам (2).
    """
    item_nums = extract_numbers(required_source_text)
    item_dates = extract_dates(required_source_text)
    item_lower = required_source_text.lower()
    item_parties = _party_hits(item_lower)
    item_doctypes = _doctype_hits(item_lower)

    results: list[tuple[Path, int, str]] = []
    for path, _f_nums, _f_dates in file_index:
        # ponytail: file_signals(path) из level0_check считает номера/даты по
        # ПОЛНОМУ пути (папка+имя слитно) — для "первое совпадение хоть где-то"
        # этого хватало, но для confident_id ниже нужно РАЗЛИЧАТЬ "номер в
        # самом имени файла" (специфичный сигнал) от "номер в имени папки"
        # (значит только "тот же договор/контур", таких файлов может быть
        # много) — поэтому считаем оба по отдельности здесь же, тем же
        # extract_numbers/extract_dates, просто на разных кусках пути.
        filename_nums = extract_numbers(path.name)
        filename_dates = extract_dates(path.name)
        folder_nums = extract_numbers(path.parent.name)

        score = 0
        why: list[str] = []

        if item_nums & filename_nums:
            score += 5
            why.append(f"номер {sorted(item_nums & filename_nums)}")

        if item_nums & folder_nums:
            score += 4
            why.append(f"номер договора в папке {sorted(item_nums & folder_nums)}")

        if item_dates & filename_dates:
            score += 3
            why.append(f"дата {sorted(item_dates & filename_dates)}")

        path_lower = str(path).lower()
        p_parties = _party_hits(path_lower)
        common_parties = item_parties & p_parties
        if common_parties:
            score += 2
            why.append(f"сторона {sorted(common_parties)[0]}")

        p_doctypes = _doctype_hits(path_lower)
        common_doctypes = item_doctypes & p_doctypes
        if common_doctypes:
            score += 2
            why.append(f"тип документа {sorted(common_doctypes)[0]}")

        if score > 0:
            results.append((path, score, "; ".join(why)))

    results.sort(key=lambda r: r[1], reverse=True)
    return results


# --- 3. исход по существу -------------------------------------------------------
#
# principle/pattern переиспользуют ladder.verify_principle_pattern (дословная
# цитата — ровно то, что можно механически сверить на уровне 0). Остальные
# card_type в ladder.py (chronology/decision/task/risk/fork) требуют полей,
# которые уровень 0 не может добыть честно и небезопасно (семантический
# «предмет решения», «стороны совпали», посчитанная сумма) — для них здесь
# генерическая сверка чисел/дат утверждения с текстом источника: то же самое
# по духу, но без притворства, что регэксп понимает предмет спора.

def verify_claim_level0(row: dict, candidates: list[tuple[Path, set[str], set[str]]]) -> dict:
    """Контракт: {"outcome": "подтверждено"|"опровергнуто"|"не смог",
    "reason": str, "evidence_path": str|None, "candidates": [пути],
    "checked": {...}}. При сомнении — «не смог» (ложное «подтверждено» —
    худший исход, 2а/приёмка)."""
    card_type = str(row.get("card_type", ""))
    claim_text = str(row.get("claim_to_verify", ""))
    items = split_sources(str(row.get("required_primary_sources", "")))
    checked: dict = {"card_type": card_type, "items": items}

    if not items:
        return {
            "outcome": "не смог",
            "reason": "в заявке не названы первоисточники (required_primary_sources пуст)",
            "evidence_path": None,
            "candidates": [],
            "checked": checked,
        }

    all_ranked = [(item, rank_candidates(item, candidates)) for item in items]

    # required_primary_sources — это "ИЛИ" из нескольких разных документов
    # (письмо; договор; переписка; ...), а claim_to_verify обычно говорит о
    # ОДНОМ конкретном из них. Если просто брать пункт с максимальным баллом
    # ранжирования, найдётся любой файл, у которого выше формальный балл —
    # например общий "договор №5239/47" вместо нужного "письмо №4722 от
    # 24.04.2026" — и заявку сверят не с тем документом. Поэтому сначала
    # пытаемся найти пункт(ы), которые сами перекликаются с числом/датой из
    # УТВЕРЖДЕНИЯ — это и есть "тот самый" источник; только если такого нет,
    # берём просто максимум по баллу.
    claim_nums_pre = extract_numbers(claim_text)
    claim_dates_pre = extract_dates(claim_text)

    def _item_relevant(item: str) -> bool:
        return bool((extract_numbers(item) & claim_nums_pre) or (extract_dates(item) & claim_dates_pre))

    relevant = [(item, ranked) for item, ranked in all_ranked if ranked and _item_relevant(item)]
    pool = relevant if relevant else all_ranked
    _, best_ranked = max(pool, key=lambda t: (t[1][0][1] if t[1] else 0))

    if not best_ranked or best_ranked[0][1] < MIN_SCORE:
        near_misses = [str(p) for _, ranked in all_ranked for p, _, _ in ranked[:3]]
        return {
            "outcome": "не смог",
            "reason": (
                f"первоисточник не опознан среди {len(candidates)} файлов 05_ORIGINALS "
                "(нет совпадения номера/даты/стороны/типа документа)"
            ),
            "evidence_path": None,
            "candidates": near_misses,
            "checked": checked,
        }

    top_path, top_score, top_why = best_ranked[0]
    cand_paths = [str(p) for p, _, _ in best_ranked[:5]]
    checked["ranking"] = top_why
    checked["ranking_score"] = top_score

    text = extract_text(top_path)
    if not text.strip():
        return {
            "outcome": "не смог",
            "reason": (
                f"источник найден ({top_path.name}), но текстовый слой пуст — "
                "нет .md/.ocr.txt рядом и нет читаемого текста в самом файле"
            ),
            "evidence_path": str(top_path),
            "candidates": cand_paths,
            "checked": checked,
        }
    checked["source_text_len"] = len(text)  # длина, не содержимое — правило контура

    # principle/pattern: дословная цитата сама себя верифицирует — случайное
    # текстовое совпадение длинной цитаты практически невозможно, поэтому
    # этой проверке можно доверять даже при слабом сигнале ранжирования
    # (в отличие от числа/даты ниже — те слишком короткие, чтобы быть
    # самодостаточным доказательством "того самого" документа).
    if card_type in ("principle", "pattern"):
        vr = ladder.verify_principle_pattern(
            card_type, {"claim_to_verify": claim_text}, {"source_excerpt": text}
        )
        checked["verifier"] = "ladder.verify_principle_pattern"
        outcome = "подтверждено" if vr["passed"] else "не смог"
        return {
            "outcome": outcome,
            "reason": vr["reason"],
            "evidence_path": str(top_path),
            "candidates": cand_paths,
            "checked": checked,
        }

    # Насколько уверенно ЭТОТ файл, а не любой другой с тем же баллом,
    # опознан как нужный источник. "номер [...]" — номер документа в самом
    # имени файла, самый специфичный сигнал. "номер договора в папке" —
    # только "тот же контур/договор", таких файлов в папке может быть
    # много — этого одного недостаточно ни для подтверждения, ни для
    # опровержения (иначе "документ упоминает номер своего же договора"
    # подтвердит почти любое утверждение об этом договоре — тавтология).
    # "сторона"/"тип документа" в одиночку — тоже неспецифично: полдюжины
    # разных исходящих писем РСУ-8 дают одинаковый балл, выбор между ними
    # был бы случайным (устойчивость sort() по порядку индексации, не по
    # смыслу). Только дата + сторона/тип вместе дают нужную специфичность.
    has_filename_number = "номер [" in top_why
    has_date = "дата [" in top_why
    has_party = "сторона" in top_why
    has_doctype = "тип документа" in top_why
    confident_id = has_filename_number or (has_date and (has_party or has_doctype))
    if not confident_id:
        return {
            "outcome": "не смог",
            "reason": (
                f"источник-кандидат {top_path.name} опознан только слабым сигналом "
                f"({top_why}) — таких файлов в 05_ORIGINALS может быть несколько, "
                "уровень 0 не доверяет содержательной сверке по такому опознанию"
            ),
            "evidence_path": None,
            "candidates": cand_paths,
            "checked": checked,
        }

    claim_nums = claim_nums_pre
    claim_dates = claim_dates_pre
    checked["claim_numbers"] = sorted(claim_nums)
    checked["claim_dates"] = sorted(claim_dates)

    if not claim_nums and not claim_dates:
        return {
            "outcome": "не смог",
            "reason": "утверждение не содержит проверяемых чисел/дат — уровню 0 сверить нечем по существу",
            "evidence_path": str(top_path),
            "candidates": cand_paths,
            "checked": checked,
        }

    doc_nums = extract_numbers(text)
    doc_dates = extract_dates(text)

    matched_nums = claim_nums & doc_nums
    matched_dates = claim_dates & doc_dates
    if matched_nums or matched_dates:
        parts = []
        if matched_nums:
            parts.append(f"номер(а) {sorted(matched_nums)}")
        if matched_dates:
            parts.append(f"дата(ы) {sorted(matched_dates)}")
        return {
            "outcome": "подтверждено",
            "reason": f"в источнике {top_path.name} найдено: {', '.join(parts)}",
            "evidence_path": str(top_path),
            "candidates": cand_paths,
            "checked": checked,
        }

    # confident_id уже гарантирован (проверено выше) — кандидат опознан
    # достаточно специфичным сигналом, несовпавшую дату/номер можно честно
    # назвать противоречием, а не «не там искали».
    if claim_dates and doc_dates and not matched_dates:
        return {
            "outcome": "опровергнуто",
            "reason": (
                f"в источнике {top_path.name} даты {sorted(doc_dates)}, "
                f"а утверждение называет {sorted(claim_dates)} — не сходится"
            ),
            "evidence_path": str(top_path),
            "candidates": cand_paths,
            "checked": checked,
        }
    if claim_nums and doc_nums and not matched_nums:
        return {
            "outcome": "опровергнуто",
            "reason": (
                f"в источнике {top_path.name} номера {sorted(doc_nums)[:5]}, "
                f"а утверждение называет {sorted(claim_nums)} — не сходится"
            ),
            "evidence_path": str(top_path),
            "candidates": cand_paths,
            "checked": checked,
        }

    return {
        "outcome": "не смог",
        "reason": (
            f"источник {top_path.name} найден и прочитан, но заявленные число/дата "
            "в тексте не встретились (текст может быть неполным/OCR-шумным)"
        ),
        "evidence_path": str(top_path),
        "candidates": cand_paths,
        "checked": checked,
    }


# --- 4. чтение реестров, ТОЛЬКО ЧТЕНИЕ ------------------------------------------

_XREF_RE = re.compile(r"(?:источники|исходники),?\s*что\s+(SV-[A-ZА-Я0-9-]+)", re.IGNORECASE)


def iter_claims():
    """Строки из всех source_verification_queue*.csv. Разрешает ссылку вида
    «те же источники, что SV-CKBA-006» на required_primary_sources ДРУГОЙ
    строки ТОГО ЖЕ прогона чтения — сама ссылка тоже только читается, ничего
    в CSV не пишется."""
    id_to_sources: dict[str, str] = {}
    rows: list[dict] = []
    for csv_path in sorted(REGISTRY_DIR.glob(REGISTRY_GLOB)):
        with csv_path.open(encoding="utf-8-sig", newline="") as fh:
            for row in csv.DictReader(fh):
                row = dict(row)
                row["registry"] = csv_path.name
                rows.append(row)
                vid = row.get("verification_id")
                if vid:
                    id_to_sources[vid] = row.get("required_primary_sources", "")

    for row in rows:
        req = row.get("required_primary_sources", "") or ""
        m = _XREF_RE.search(req)
        if m and m.group(1) in id_to_sources:
            row = dict(row)
            row["required_primary_sources"] = id_to_sources[m.group(1)]
        yield row


# --- 5. CLI: --dry-run -----------------------------------------------------------

def _dry_run() -> None:
    file_index = load_originals(ORIGINALS_DIR)
    outcomes = {"подтверждено": 0, "опровергнуто": 0, "не смог": 0}
    with_candidates = 0
    by_type: dict[str, dict[str, int]] = {}
    rows_out: list[dict] = []

    for row in iter_claims():
        result = verify_claim_level0(row, file_index)
        outcomes[result["outcome"]] += 1
        if result["candidates"]:
            with_candidates += 1
        ct = row.get("card_type", "")
        bucket = by_type.setdefault(ct, {"подтверждено": 0, "опровергнуто": 0, "не смог": 0})
        bucket[result["outcome"]] += 1
        rows_out.append(
            {
                "verification_id": row.get("verification_id"),
                "registry": row.get("registry"),
                "card_type": ct,
                "outcome": result["outcome"],
                "reason": result["reason"],
                "evidence_path": result["evidence_path"],
                "n_candidates": len(result["candidates"]),
            }
        )

    summary = {
        "total": len(rows_out),
        "outcomes": outcomes,
        "rows_with_candidates": with_candidates,
        "by_card_type": by_type,
        "files_indexed_in_05_ORIGINALS": len(file_index),
    }

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        json.dumps({"summary": summary, "rows": rows_out}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    print(f"\nreport written to: {REPORT_PATH}")


# --- 6. самопроверка --------------------------------------------------------------

def _selftest() -> None:
    # 1. транслитерация сторон — совпадает в обе стороны
    assert "rsu8" in _party_hits("исходящее письмо рсу-8 №15")
    assert "ckba" in _party_hits("письмо цкба №47")
    assert _party_hits("рсу-8") & _party_hits("rsu8-out-file.pdf")
    assert not _party_hits("переписка сторон")  # нет узнаваемой стороны

    # 2. ранжирование: правильный кандидат выше мусора
    p_letter = Path(r"E:\-5-\030_CKBA_Wiki\05_ORIGINALS\CHT-020\RSU8-OUT-2026-07-22-suspension-VG-act.docx")
    p_noise = Path(r"E:\-5-\030_CKBA_Wiki\05_ORIGINALS\XLSX\smeta_customers.xlsx")
    idx = [(p, *file_signals(p)) for p in (p_letter, p_noise)]
    ranked = rank_candidates("исходящее письмо РСУ-8 от 22.07.2026, акт", idx)
    assert ranked and ranked[0][0] == p_letter, "ожидали письмо РСУ-8 первым кандидатом"
    assert ranked[0][1] >= MIN_SCORE

    # 3. без опознаваемых сигналов — пусто (нечего ранжировать)
    assert rank_candidates("переписка сторон", idx) == []

    # 4. извлечение текста из .md-слоя (реальный сайдкар, без PDF-библиотек)
    p_pdf_with_sidecar = Path(r"E:\-5-\030_CKBA_Wiki\05_ORIGINALS\CHT-020\8564_suspension.pdf")
    text = extract_text(p_pdf_with_sidecar)
    assert text.strip(), ".md-слой должен был дать непустой текст"
    assert "5239" in extract_numbers(text)  # номер договора реально есть в тексте письма

    # 5. verify_claim_level0: «подтверждено» — номер из утверждения найден в реальном тексте
    row_ok = {
        "card_type": "decision",
        "claim_to_verify": "Работы по договору №5239/47 приостановлены",
        "required_primary_sources": "исходящее письмо РСУ-8 от 22.07.2026, акт о приостановке №5239/47",
    }
    res_ok = verify_claim_level0(row_ok, [(p_letter, *file_signals(p_letter))])
    assert res_ok["outcome"] == "подтверждено", res_ok["reason"]
    assert res_ok["evidence_path"] == str(p_letter)

    # 6. verify_claim_level0: при сомнении — «не смог» (кандидат опознан только
    #    по дате — не независимый сигнал, — заявленный номер не встретился)
    row_unsure = {
        "card_type": "task",
        "claim_to_verify": "Документ содержит номер 999888",
        "required_primary_sources": "документ от 22.07.2026",
    }
    res_unsure = verify_claim_level0(row_unsure, [(p_letter, *file_signals(p_letter))])
    assert res_unsure["outcome"] == "не смог", res_unsure["reason"]

    # 7. verify_claim_level0: «опровергнуто» — кандидат опознан УВЕРЕННО
    #    (дата+сторона+тип документа сошлись при ранжировании), а заявленная
    #    в утверждении дата противоречит датам В ТЕКСТЕ источника
    row_bad_date = {
        "card_type": "decision",
        "claim_to_verify": "Акт составлен 01.01.2020",
        "required_primary_sources": "исходящее письмо РСУ-8 от 22.07.2026, акт о приостановке",
    }
    res_bad = verify_claim_level0(row_bad_date, [(p_letter, *file_signals(p_letter))])
    assert res_bad["outcome"] == "опровергнуто", res_bad["reason"]

    # 7б. тот же кандидат, но опознан ТОЛЬКО слабым сигналом (сторона+тип
    #     документа, без даты) — таких писем РСУ-8 может быть много, поэтому
    #     даже настоящее несовпадение даты не эскалируется в "опровергнуто"
    row_weak_id = {
        "card_type": "decision",
        "claim_to_verify": "Акт составлен 01.01.2020",
        "required_primary_sources": "исходящее письмо РСУ-8, акт о приостановке",
    }
    res_weak = verify_claim_level0(row_weak_id, [(p_letter, *file_signals(p_letter))])
    assert res_weak["outcome"] == "не смог", res_weak["reason"]

    # 8. principle/pattern через ladder.verify_principle_pattern (реюз, не копия)
    row_quote_bad = {
        "card_type": "principle",
        "claim_to_verify": "цитата, которой в документе точно нет",
        "required_primary_sources": "исходящее письмо РСУ-8, акт о приостановке",
    }
    res_quote = verify_claim_level0(row_quote_bad, [(p_letter, *file_signals(p_letter))])
    assert res_quote["outcome"] == "не смог"  # цитата не найдена дословно — не "опровергнуто"

    # 9. нет первоисточников в заявке — честное "не смог", без исключений
    res_empty = verify_claim_level0(
        {"card_type": "risk", "claim_to_verify": "x", "required_primary_sources": ""}, idx
    )
    assert res_empty["outcome"] == "не смог"
    assert res_empty["candidates"] == []

    # 10. кэш: второй вызов extract_text на некэшируемом напрямую .xlsx-файле
    #     (без .md-слоя) отдаёт то же самое и не падает
    p_xlsx = Path(r"E:\-5-\030_CKBA_Wiki\05_ORIGINALS\XLSX\smeta_customers.xlsx")
    if p_xlsx.is_file():
        t1 = extract_text(p_xlsx)
        t2 = extract_text(p_xlsx)
        assert t1 == t2
        assert _cache_path(p_xlsx).is_file()

    # 11. неподдержанные форматы — пустая строка, не исключение
    assert extract_text(Path(r"E:\-5-\030_CKBA_Wiki\05_ORIGINALS\CHT-020\7. Еще договора ЦКБА.zip")) == ""

    print("selftest ok")


if __name__ == "__main__":
    if "--selftest" in sys.argv:
        _selftest()
    elif "--dry-run" in sys.argv:
        _dry_run()
    else:
        print(__doc__)
